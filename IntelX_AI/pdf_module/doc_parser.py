"""
DOC/DOCX file parser module.
Extracts data from Word documents and converts to text format for LLM processing.
"""
from pathlib import Path
from typing import Dict, Any, List
import sys
import os

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from exceptions import DOCParserError, FileHandlerError, DependencyError


def parse_docx_file(docx_path: str) -> str:
    """
    Parse DOCX file and convert to text format for LLM extraction.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        String representation of document content
    """
    if not docx_path:
        raise DOCParserError("DOCX file path is required", docx_path)
    
    try:
        from docx import Document
    except ImportError:
        raise DependencyError(
            "python-docx is required to read DOCX files. Install it with `pip install python-docx`.",
            docx_path
        )
    
    path = Path(docx_path)
    
    if not path.exists():
        raise FileHandlerError(f"DOCX file not found: {docx_path}", docx_path)
    
    if not path.is_file():
        raise FileHandlerError(f"Path is not a file: {docx_path}", docx_path)
    
    try:
        document = Document(path)
        
        text_parts = []
        text_parts.append(f"Document: {path.name}")
        text_parts.append("=" * 50)
        text_parts.append("")
        
        # Extract paragraphs
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        
        if paragraphs:
            text_parts.append("Content:")
            text_parts.append("")
            for idx, para in enumerate(paragraphs, 1):
                text_parts.append(f"{para}")
                text_parts.append("")
        
        # Extract tables
        if document.tables:
            text_parts.append("")
            text_parts.append("Tables:")
            text_parts.append("=" * 50)
            
            for table_idx, table in enumerate(document.tables, 1):
                text_parts.append(f"\nTable {table_idx}:")
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace('\n', ' ')
                        row_data.append(cell_text)
                    text_parts.append(" | ".join(row_data))
                
                text_parts.append("")
        
        return "\n".join(text_parts)
    
    except FileNotFoundError as e:
        raise FileHandlerError(f"DOCX file not found: {docx_path}", docx_path) from e
    except PermissionError as e:
        raise FileHandlerError(f"Permission denied reading DOCX file: {docx_path}", docx_path) from e
    except Exception as e:
        if "docx" in str(type(e).__module__).lower():
            raise DOCParserError(f"python-docx error: {str(e)}", docx_path) from e
        raise DOCParserError(f"Unexpected error parsing DOCX file: {str(e)}", docx_path) from e


def parse_doc_file(doc_path: str) -> str:
    """
    Parse DOC file (older Word format) and convert to text format.
    Note: This requires additional libraries like textract or antiword.
    
    Args:
        doc_path: Path to the DOC file
        
    Returns:
        String representation of document content
        
    Raises:
        DOCParserError: If parsing fails
        FileHandlerError: If file operations fail
        DependencyError: If required tools are missing
    """
    if not doc_path:
        raise DOCParserError("DOC file path is required", doc_path)
    
    path = Path(doc_path)
    
    if not path.exists():
        raise FileHandlerError(f"DOC file not found: {doc_path}", doc_path)
    
    try:
        # Try using textract if available
        import textract
        try:
            text = textract.process(str(path)).decode('utf-8')
            return f"Document: {path.name}\n{'=' * 50}\n\n{text}"
        except Exception as e:
            raise DOCParserError(f"textract failed to extract text: {str(e)}", doc_path) from e
    except ImportError:
        try:
            # Try using antiword if available (Linux)
            import subprocess
            result = subprocess.run(
                ['antiword', str(path)],
                capture_output=True,
                text=True,
                encoding='utf-8',
                timeout=30
            )
            if result.returncode == 0:
                return f"Document: {path.name}\n{'=' * 50}\n\n{result.stdout}"
            else:
                raise DOCParserError(
                    f"antiword failed to extract text. Error: {result.stderr}",
                    doc_path
                )
        except FileNotFoundError:
            raise DependencyError(
                "To parse DOC files, install either 'textract' (pip install textract) or 'antiword' (apt-get install antiword)",
                doc_path
            )
        except subprocess.TimeoutExpired:
            raise DOCParserError("antiword process timed out", doc_path)
    except Exception as e:
        raise DOCParserError(f"Unexpected error parsing DOC file: {str(e)}", doc_path) from e


def parse_doc_file_generic(doc_path: str) -> str:
    """
    Generic DOC/DOCX parser that handles both formats.
    
    Args:
        doc_path: Path to the DOC or DOCX file
        
    Returns:
        String representation of document content
        
    Raises:
        DOCParserError: If parsing fails
        FileHandlerError: If file operations fail
    """
    if not doc_path:
        raise DOCParserError("Document file path is required", doc_path)
    
    path = Path(doc_path)
    
    if not path.exists():
        raise FileHandlerError(f"Document file not found: {doc_path}", doc_path)
    
    suffix = path.suffix.lower()
    
    if suffix == '.docx':
        return parse_docx_file(doc_path)
    elif suffix == '.doc':
        return parse_doc_file(doc_path)
    else:
        raise DOCParserError(f"Unsupported file format: {suffix}. Supported formats: .doc, .docx", doc_path)


def extract_doc_data(doc_path: str) -> Dict[str, Any]:
    """
    Extract structured data from DOC/DOCX file.
    
    Args:
        doc_path: Path to the DOC or DOCX file
        
    Returns:
        Dictionary with document structure and data
        
    Raises:
        DOCParserError: If parsing fails
        FileHandlerError: If file operations fail
        DependencyError: If required dependencies are missing
    """
    if not doc_path:
        raise DOCParserError("Document file path is required", doc_path)
    
    path = Path(doc_path)
    
    if not path.exists():
        raise FileHandlerError(f"Document file not found: {doc_path}", doc_path)
    
    suffix = path.suffix.lower()
    
    if suffix == '.docx':
        try:
            from docx import Document
        except ImportError:
            raise DependencyError(
                "python-docx is required to read DOCX files. Install it with `pip install python-docx`.",
                doc_path
            )
        
        try:
            document = Document(path)
            
            paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            
            tables_data = []
            for table_idx, table in enumerate(document.tables, 1):
                table_rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_data)
                tables_data.append({
                    "table_number": table_idx,
                    "rows": table_rows,
                    "row_count": len(table_rows),
                    "column_count": len(table_rows[0]) if table_rows else 0
                })
            
            return {
                "file_path": str(path),
                "file_type": "docx",
                "paragraphs": paragraphs,
                "paragraph_count": len(paragraphs),
                "tables": tables_data,
                "table_count": len(tables_data)
            }
        except Exception as e:
            raise DOCParserError(f"Error extracting DOCX data: {str(e)}", doc_path) from e
    
    elif suffix == '.doc':
        # For DOC files, return basic info since parsing requires external tools
        return {
            "file_path": str(path),
            "file_type": "doc",
            "note": "DOC file parsing requires textract or antiword. Use parse_doc_file() for text extraction.",
            "paragraphs": [],
            "tables": []
        }
    else:
        raise DOCParserError(f"Unsupported file format: {suffix}. Supported formats: .doc, .docx", doc_path)

